"""
Generate paper publication DOCX in Sandip University college letterhead format.
Logo is fetched from URL at runtime — no local file needed.
"""

import io
import urllib.request
import tempfile
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Sandip University branding ───────────────────────────────────────────────
LOGO_URL = "https://ik.imagekit.io/fe8cguiig/sandip%20university%20logo.png?updatedAt=1773121773871"
UNIVERSITY_NAME    = "Sandip University, Nashik (MS), India"
UNIVERSITY_ADDRESS = "At Post Mahiravani, Trimbak Road, Nashik-422213, Maharashtra"
UNIVERSITY_WEB     = "https://www.sandipuniversity.edu.in"


def fetch_logo_to_tempfile():
    """Download logo from URL and save to a temp file. Returns path or None."""
    try:
        req = urllib.request.Request(
            LOGO_URL,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
                "Accept": "image/png,image/*,*/*",
                "Accept-Language": "en-US,en;q=0.9",
            }
        )
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = resp.read()
        if len(data) < 100:   # sanity check — not an empty/error response
            return None
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        tmp.write(data)
        tmp.close()
        return tmp.name
    except Exception:
        return None


# ── Helpers ───────────────────────────────────────────────────────────────────
def remove_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def shade_cell(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


def set_row_height(row, twips, rule='atLeast'):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for existing in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(twips))
    trHeight.set(qn('w:hRule'), rule)
    trPr.append(trHeight)


def vmerge_start(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vMerge = OxmlElement('w:vMerge')
    vMerge.set(qn('w:val'), 'restart')
    tcPr.append(vMerge)


def vmerge_continue(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vMerge = OxmlElement('w:vMerge')
    tcPr.append(vMerge)
    for p in tc.findall('.//' + qn('w:p')):
        for r in p.findall(qn('w:r')):
            p.remove(r)


def cell_para(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
              color=None, italic=False, space_before=0, space_after=0):
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


# ── Main generator ────────────────────────────────────────────────────────────
def generate_paper_publication_doc(papers, stage, logo_path=None):
    """
    Generate DOCX for paper publications in Sandip University letterhead format.

    papers: list of dicts with keys:
        sr_no, paper_title, journal_name, volume_no, issue, timeline, e_issn,
        prn_list (list of up to 4 strings), name_list (list of up to 4 strings)
    stage: 1 or 2
    logo_path: (ignored — logo is fetched from URL automatically)
    """
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.5)
    section.header_distance = Cm(0.5)

    # ── Fetch logo ────────────────────────────────────────────────────────
    tmp_logo = fetch_logo_to_tempfile()

    # ── HEADER: Logo in page header area ─────────────────────────────────
    header = section.header
    # Clear default empty paragraph
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    h_para = header.add_paragraph()
    h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_para.paragraph_format.space_before = Pt(0)
    h_para.paragraph_format.space_after  = Pt(0)

    logo_inserted = False
    if tmp_logo:
        try:
            run = h_para.add_run()
            run.add_picture(tmp_logo, width=Inches(5.6))
            logo_inserted = True
        except Exception:
            pass

    if not logo_inserted:
        # Fallback: draw university name as large bold text in the header
        h_para.clear()
        h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = h_para.add_run(UNIVERSITY_NAME)
        r1.bold = True
        r1.font.size = Pt(15)
        r1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        h_para2 = header.add_paragraph()
        h_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h_para2.paragraph_format.space_before = Pt(0)
        h_para2.paragraph_format.space_after  = Pt(0)
        r2 = h_para2.add_run(UNIVERSITY_ADDRESS)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        h_para3 = header.add_paragraph()
        h_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h_para3.paragraph_format.space_before = Pt(0)
        h_para3.paragraph_format.space_after  = Pt(0)
        r3 = h_para3.add_run(UNIVERSITY_WEB)
        r3.font.size = Pt(10)
        r3.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ── University name block (body, below header) ────────────────────────
    def center_para(text, size=11, bold=False, space_after=2, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        return p

    # Only show university text in body if logo was inserted in header.
    # When logo fetch fails, the fallback header already has the university
    # info — skip here to avoid showing it twice.
    if logo_inserted:
        center_para(UNIVERSITY_NAME,    size=13, bold=True,  space_after=1, color=RGBColor(0x1A, 0x1A, 0x2E))
        center_para(UNIVERSITY_ADDRESS, size=10, bold=False, space_after=1, color=RGBColor(0x1A, 0x1A, 0x2E))
        center_para(UNIVERSITY_WEB,     size=10, bold=False, space_after=4, color=RGBColor(0x1A, 0x1A, 0x2E))

    # Thin divider line under university block
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_before = Pt(0)
    div_p.paragraph_format.space_after  = Pt(6)
    pPr = div_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Department info
    center_para("School of Computer Science and Engineering", size=12, bold=True,  space_after=1)
    center_para("Department of Computer Science and Engineering", size=11, bold=False, space_after=1)
    center_para("Academic Year 2023-24", size=11, bold=False, space_after=6)

    # ── MAIN TABLE ────────────────────────────────────────────────────────
    # Columns: SR.NO. | PAPER TITLE | PRN | NAME OF STUDENT
    # Content width = 21 - 1.8 - 1.5 = 17.7 cm
    COL_WIDTHS = [Cm(1.3), Cm(7.5), Cm(3.0), Cm(5.9)]

    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # ── Row 0: Title banner ───────────────────────────────────────────────
    r0 = table.add_row()
    set_row_height(r0, 520)
    merged = r0.cells[0].merge(r0.cells[3])
    shade_cell(merged, "1F3864")
    merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_para(merged,
              f"Students Paper Publications — Stage {stage}",
              bold=True, size=13,
              color=RGBColor(255, 255, 255),
              align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before=4, space_after=4)

    # ── Row 1: Class / Date ───────────────────────────────────────────────
    r1 = table.add_row()
    set_row_height(r1, 360)
    left  = r1.cells[0].merge(r1.cells[1])
    right = r1.cells[2].merge(r1.cells[3])
    shade_cell(left,  "D6E4F0")
    shade_cell(right, "D6E4F0")
    left.vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_para(left,  "Class: B.Tech CSE",       size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT,  space_before=2, space_after=2)
    cell_para(right, "Date: _________________", size=10.5, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=2, space_after=2)

    # ── Row 2: Column headers ─────────────────────────────────────────────
    r2 = table.add_row()
    set_row_height(r2, 480)
    hdrs = ["SR.\nNO.", "PAPER TITLE", "PRN", "NAME OF STUDENT"]
    for ci, (cell, htext) in enumerate(zip(r2.cells, hdrs)):
        cell.width = COL_WIDTHS[ci]
        shade_cell(cell, "2E5596")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_para(cell, htext, bold=True, size=10,
                  color=RGBColor(255, 255, 255),
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=3, space_after=3)

    # ── Data rows ─────────────────────────────────────────────────────────
    for idx, paper in enumerate(papers):
        fill    = "EEF3FA" if idx % 2 == 0 else "FFFFFF"
        prns    = (paper.get('prn_list') or [])
        names   = (paper.get('name_list') or [])
        n_members = max(1, len([x for x in names if x]))

        citation = (
            f'PUBLISHED PAPER ENTITLED "{paper["paper_title"].upper()}" '
            f'IN {paper["journal_name"].upper()} '
            f'VOLUME {paper["volume_no"]}, ISSUE {paper["issue"]}, '
            f'{paper["timeline"].upper()}, E ISSN {paper["e_issn"]}'
        )

        for mi in range(n_members):
            dr = table.add_row()
            set_row_height(dr, 420, 'atLeast')
            for ci in range(4):
                dr.cells[ci].width = COL_WIDTHS[ci]
                shade_cell(dr.cells[ci], fill)
                dr.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if mi == 0:
                vmerge_start(dr.cells[0])
                cell_para(dr.cells[0], str(paper['sr_no']),
                          bold=True, size=11,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
                vmerge_start(dr.cells[1])
                cell_para(dr.cells[1], citation,
                          size=8.5, italic=True,
                          align=WD_ALIGN_PARAGRAPH.LEFT,
                          space_before=2, space_after=2)
            else:
                vmerge_continue(dr.cells[0])
                vmerge_continue(dr.cells[1])

            prn  = prns[mi]  if mi < len(prns)  else ''
            name = names[mi] if mi < len(names) else ''
            cell_para(dr.cells[2], prn,  size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_para(dr.cells[3], name, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Signature section ─────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_before = Pt(20)

    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in sig_tbl.rows:
        for cell in row.cells:
            remove_borders(cell)

    lc = sig_tbl.rows[0].cells[0]
    lc.width = Cm(8.5)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(0)
    lp.add_run("\n\n\n\nProject Coordinator").bold = True

    rc = sig_tbl.rows[0].cells[1]
    rc.width = Cm(9.2)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after  = Pt(0)
    rp.add_run("\n\n\n\nHOD").bold = True

    # ── Copy distribution ─────────────────────────────────────────────────
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(16)
    cp.paragraph_format.space_after  = Pt(0)
    r = cp.add_run("Copy(s) to:")
    r.bold = True
    r.font.size = Pt(10)

    for item in ["1. Project File", "2. Student Publication File"]:
        ip = doc.add_paragraph()
        ip.paragraph_format.space_before = Pt(0)
        ip.paragraph_format.space_after  = Pt(0)
        ip.add_run(item).font.size = Pt(10)

    # ── Save & cleanup ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    if tmp_logo and os.path.exists(tmp_logo):
        os.remove(tmp_logo)

    return buf


# ── Standalone test ───────────────────────────────────────────────────────────
if __name__ == '__main__':
    sample = [
        {
            'sr_no': 1,
            'paper_title': 'AI Based Face Recognition Attendance System',
            'journal_name': 'IJIRMPS',
            'volume_no': '12', 'issue': '2',
            'timeline': 'March-April 2024',
            'e_issn': '2349-7300',
            'prn_list':  ['22CE001', '22CE002', '22CE003', '22CE004'],
            'name_list': ['Arun Kumar', 'Priya Sharma', 'Raj Patel', 'Sneha Iyer'],
        },
        {
            'sr_no': 2,
            'paper_title': 'Smart Traffic Management Using IoT and Deep Learning',
            'journal_name': 'IJRASET',
            'volume_no': '11', 'issue': '5',
            'timeline': 'May 2024',
            'e_issn': '2321-9653',
            'prn_list':  ['22CE005', '22CE006'],
            'name_list': ['Kiran Singh', 'Meera Nair'],
        },
    ]
    print("Fetching logo from URL...")
    buf = generate_paper_publication_doc(sample, stage=1)
    with open('sandip_test_output.docx', 'wb') as f:
        f.write(buf.read())
    print("Generated sandip_test_output.docx")